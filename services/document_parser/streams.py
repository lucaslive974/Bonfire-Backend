import io
import re
import unicodedata
from typing import Any, Dict, Generator, List

import numpy as np
import pandas as pd
from docx import Document
from docx.table import Table
from pyingestion import ExtractionSession, InputStream, OutputStream, TransformStream

from classes.Conversores import Conversores
from exceptions.CustomExceptions import (
    ErrDataPubli,
    ErrIncorrectInstance,
    ErrQuantityOfAtas,
    ErrReadingFile,
)
from services.document_parser.pubsub import AsyncMessageProcessor
from services.parsers import normalize_auto_infraction_id
from services.recurso_service import RecursoService

# ==========================================
# WRITE STREAMS (Output)
# ==========================================


class BonfireRecursoWriteStream(OutputStream[Any]):
    def __init__(self, first_instance: bool = True, batch_size: int = 100):
        self.first_instance = first_instance
        self.processor = AsyncMessageProcessor(
            self._process_batch, batch_size=batch_size
        )
        self.processor.start()
        super().__init__()

    def write(self, item: Any) -> None:
        if isinstance(item, list):
            for i in item:
                self.processor.publish(i)
        else:
            self.processor.publish(item)

    def _process_batch(self, batch: List[Any]) -> None:
        if self.first_instance:
            RecursoService.insert_primeira_instancia(batch)
        else:
            RecursoService.insert_segunda_instancia(batch)

    def flush(self) -> None:
        self.processor.stop()


class BonfireInfracaoWriteStream(OutputStream[Any]):
    def __init__(self, ignore: bool = False, batch_size: int = 1000):
        self.ignore = ignore
        self.processor = AsyncMessageProcessor(
            self._process_batch, batch_size=batch_size
        )
        self.processor.start()
        super().__init__()

    def write(self, item: Any) -> None:
        if isinstance(item, list):
            for i in item:
                self.processor.publish(i)
        else:
            self.processor.publish(item)

    def _process_batch(self, batch: List[Any]) -> None:
        from repositories.autoinfracao_repository import AutoInfracaoRepository
        from repositories.database import get_db

        with get_db() as db:
            repo = AutoInfracaoRepository(db)
            repo.insert_bulk_rows(batch, ignore=self.ignore)

    def flush(self) -> None:
        self.processor.stop()


# ==========================================
# INPUT STREAMS
# ==========================================


class RecursosDocxInputStream(InputStream[Any, Dict[str, Any]]):
    """
    Substitui a lógica de parse_docx_recursos, extraindo recursos do DOCX
    e fazendo o yield item por item, respeitando o modelo de stream.
    """

    def __init__(self, first_instance: bool = True):
        self.first_instance = first_instance
        self.current_unit_index = 0
        self.total_units = 0

    def extract_data_publ(self, doc: Document) -> str:
        data_publicacao_extracted = " ".join([p.text for p in doc.paragraphs])
        padrao_data = r"PUBLICADO NO DI[ÁA]RIO OFICIAL DO MUNIC[ÍI]PIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})"
        match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)
        dat_publ = match_data_publicacao.group(1) if match_data_publicacao else None

        if dat_publ is not None:
            dat_publ = Conversores.converte_data(dat_publ)
        else:
            raise ErrDataPubli("Data de publicação não encontrada no documento", 400)
        return dat_publ

    def extract_atas(self, doc: Document) -> list[int]:
        num_atas = []
        padrao_num_ata = r"ATA\s+DA\s+(\d+)ª"
        for paragraph in doc.paragraphs:
            match_num_ata = re.search(padrao_num_ata, paragraph.text)
            if match_num_ata:
                num_atas.append(match_num_ata.group(1))

        qtd_atas = len(num_atas)
        qtd_tables = len(doc.tables)
        if not self.first_instance and len(num_atas) > 0:
            raise ErrIncorrectInstance(
                "Instância incorreta. Importe como recurso de primeira instância"
            )
        if self.first_instance and (qtd_atas != qtd_tables):
            raise ErrQuantityOfAtas(
                "Quantidade de atas encontradas difere da quantidade de tabelas",
                qtd_atas,
                qtd_tables,
                400,
            )
        return num_atas or None

    def process_table(self, table: Table, dat_publ: str, num_ata: int | None):
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            num_recurso = row_data[0]
            if num_recurso == "RECURSO":
                continue
            num_ai = normalize_auto_infraction_id(row_data[1])
            nom_conc = row_data[2]
            valida_resultado = str(row_data[3]).upper()
            resultado = valida_resultado != "IMPROCEDENTE"

            recurso = {
                "NUM_RECURSO": num_recurso,
                "NUM_AI": num_ai,
                "NOM_CONC": nom_conc,
                "RESULTADO": resultado,
                "DAT_PUBL": dat_publ,
            }
            if num_ata:
                recurso["NUM_ATA"] = num_ata
            yield recurso

    def read(
        self, source: Any, session: ExtractionSession | None = None
    ) -> Generator[Dict[str, Any], None, None]:
        doc = Document(source)
        dat_publ = self.extract_data_publ(doc)
        num_atas = self.extract_atas(doc)

        for index, table in enumerate(doc.tables):
            num_ata = num_atas[index] if self.first_instance else None
            if session:
                session.process_page_result(True, 1, 1)
            yield from self.process_table(table, dat_publ, num_ata)


class SanitizedTextIO(io.TextIOBase):
    """
    Wrapper that intercepts the binary file stream,
    removes accents and encoding garbage BEFORE pandas attempts to parse it.
    """

    def __init__(self, binary_stream):
        self.binary_stream = binary_stream

    def read(self, size=-1):
        raw_bytes = self.binary_stream.read(size)
        if not raw_bytes:
            return ""

        # 1. Try UTF-8 or Latin-1, forcing the drop (ignore) of corrupted bytes
        try:
            text = raw_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = raw_bytes.decode("latin_1", errors="ignore")

        # 2. Normalize NFKD (remove accents) and drop any non-ASCII character
        return (
            unicodedata.normalize("NFKD", text)
            .encode("ascii", "ignore")
            .decode("ascii")
        )

    def seek(self, offset, whence=0):
        return self.binary_stream.seek(offset, whence)


class InfracoesCsvInputStream(InputStream[Any, pd.DataFrame]):
    """InputStream for CSV infraction files."""

    def __init__(self):
        self.current_unit_index = 0
        self.total_units = 0

    def _detect_separator(self, source: Any) -> str:
        """Reads the first bytes of the stream to infer the separator and resets the cursor."""
        try:
            first_chunk = source.read(2048)
            try:
                text = first_chunk.decode("utf-8", errors="ignore")
            except Exception:
                text = first_chunk.decode("latin_1", errors="ignore")

            sep_counts = {
                ",": text.count(","),
                ";": text.count(";"),
                "|": text.count("|"),
            }
            best_sep = max(sep_counts, key=sep_counts.get)

            source.seek(0)
            return best_sep if sep_counts[best_sep] > 0 else ";"
        except Exception:
            try:
                source.seek(0)
            except Exception:
                pass
            return ";"

    def read(
        self, source: Any, session: ExtractionSession | None = None
    ) -> Generator[pd.DataFrame, None, None]:
        try:
            best_sep = self._detect_separator(source)
            clean_source = SanitizedTextIO(source)
            for chunk in pd.read_csv(
                clean_source, header=0, delimiter=best_sep, chunksize=1000
            ):
                yield chunk
        except Exception as e:
            raise ErrReadingFile(f"Severe failure reading CSV: {str(e)}", 500)


class InfracoesXlsInputStream(InputStream[Any, pd.DataFrame]):
    """InputStream para arquivos de infração XLS."""

    def __init__(self):
        self.current_unit_index = 0
        self.total_units = 0

    def read(
        self, source: Any, session: ExtractionSession | None = None
    ) -> Generator[pd.DataFrame, None, None]:
        try:
            data_frame = pd.read_excel(source, header=0)
            yield data_frame
        except Exception as e:
            raise ErrReadingFile(
                f"Problema ao processar o arquivo no Load: {source}. {e}", 500
            )


# ==========================================
# TRANSFORM STREAMS
# ==========================================


class InfracoesTransformStream(TransformStream[pd.DataFrame, List[Dict[str, Any]]]):
    """
    Transforma o DataFrame de Infrações formatando as colunas e datas.
    Recebe os formatos de data/hora no construtor para reaproveitar a lógica entre CSV e XLS.
    """

    def __init__(
        self, datetime_format: str, date_format: str, convert_val_infr: bool = False
    ):
        self.datetime_format = datetime_format
        self.date_format = date_format
        self.convert_val_infr = convert_val_infr
        super().__init__()

    def transform(self, data_frame: pd.DataFrame) -> List[Dict[str, Any]]:
        try:
            if "HORA" in data_frame.columns:
                data_frame["DAT_OCOR_INFR"] = (
                    data_frame["DAT_OCOR_INFR"].astype(str)
                    + " "
                    + data_frame["HORA"].astype(str)
                )

            data_frame["DAT_OCOR_INFR"] = pd.to_datetime(
                data_frame["DAT_OCOR_INFR"], format=self.datetime_format
            )
            data_frame["DAT_EMIS_NOTF"] = pd.to_datetime(
                data_frame["DAT_EMIS_NOTF"], format=self.date_format
            )
            data_frame["DAT_LIMT_RECU"] = pd.to_datetime(
                data_frame["DAT_LIMT_RECU"], format=self.date_format
            )

            if self.convert_val_infr and "VAL_INFR" in data_frame.columns:
                data_frame["VAL_INFR"] = data_frame["VAL_INFR"].map(
                    lambda x: pd.to_numeric(str(x).replace(",", "."))
                )

            if "DAT_CANC" in data_frame.columns and not bool(
                data_frame["DAT_CANC"].isnull().all()
            ):
                data_frame["DAT_CANC"] = pd.to_datetime(
                    data_frame["DAT_CANC"], format=self.date_format
                )

            if "HORA" in data_frame.columns:
                data_frame = data_frame.drop(columns=["HORA"])

            data_frame.replace([np.nan], [None], inplace=True)
            return data_frame.to_dict(orient="records")
        except Exception as e:
            raise ErrReadingFile(f"Erro no transform de Infrações. {e}", 500)


class NoOpTransformStream(TransformStream[Any, Any]):
    """Transform Stream que não altera os dados, apenas passa adiante."""

    def transform(self, data: Any) -> Any:
        return data
