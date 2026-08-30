import re
from typing import Any, Dict, List

import numpy as np
import pandas as pd
from docx import Document

from classes.Conversores import Conversores
from exceptions.CustomExceptions import (
    ErrDataPubli,
    ErrIncorrectInstance,
    ErrQuantityOfAtas,
    ErrReadingFile,
)


def normalize_auto_infraction_id(ai: str) -> str:
    if not ai:
        return ai

    if "-" in ai:
        return ai

    if ai[len(ai) - 1] == "-":
        return ai

    return f"{ai[:-1]}-{ai[-1]}"


def parse_docx_recursos(docx: Any, first_instance: bool = True) -> List[Dict[str, Any]]:
    """Realiza o parse do documento DOCX de resultado de recurso e extrai os recursos como dicionários"""
    doc = Document(docx)
    recurso_list = []

    # Extraindo data da publicação
    data_publicacao_extracted = ""
    for paragraph in doc.paragraphs:
        data_publicacao_extracted += paragraph.text + " "

    padrao_data = r'PUBLICADO NO DI[ÁA]RIO OFICIAL DO MUNIC[ÍI]PIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})'
    match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)

    dat_publ = match_data_publicacao.group(1) if match_data_publicacao else None

    if dat_publ is not None:
        dat_publ = Conversores.converte_data(dat_publ)
    else:
        raise ErrDataPubli("Data de publicação não encontrada no documento", 400)

    # Extraindo Numero da ata
    num_atas = []
    padrao_num_ata = r'ATA\s+DA\s+(\d+)ª'
    for paragraph in doc.paragraphs:
        match_num_ata = re.search(padrao_num_ata, paragraph.text)
        if match_num_ata:
            num_atas.append(match_num_ata.group(1))

    qtd_atas = len(num_atas)
    qtd_tables = len(doc.tables)
    if (qtd_atas != qtd_tables) and first_instance:
        raise ErrQuantityOfAtas("Quantidade de atas encontradas difere da quantidade de tabelas", qtd_atas, qtd_tables, 400)

    for index, table in enumerate(doc.tables):
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            num_recurso = row_data[0]
            num_ai = normalize_auto_infraction_id(row_data[1])
            nom_conc = row_data[2]
            valida_resultado = str(row_data[3]).upper()
            if valida_resultado == 'IMPROCEDENTE':
                resultado = False
            else:
                resultado = True

            if not first_instance and num_atas[index] is not None:
                raise ErrIncorrectInstance("Instância incorreta. Importe como recurso de primeira instância")

            num_ata = num_atas[index] if first_instance else 0

            if num_recurso == 'RECURSO':
                continue

            recurso_dict = {
                'NUM_RECURSO': num_recurso,
                'NUM_ATA': num_ata,
                'NUM_AI': num_ai,
                'NOM_CONC': nom_conc,
                'RESULTADO': resultado,
                'DAT_PUBL': dat_publ
            }
            recurso_list.append(recurso_dict)
    return recurso_list


def parse_csv_infracoes(csv: Any) -> pd.DataFrame:
    """Processa o arquivo CSV de infrações retornando um DataFrame formatado"""
    try:
        data_frame = pd.read_csv(csv, header=0, encoding="latin_1", delimiter=';')
        data_frame['DAT_OCOR_INFR'] = data_frame['DAT_OCOR_INFR'].astype(str) + " " + data_frame['HORA'].astype(str)
        data_frame['DAT_OCOR_INFR'] = pd.to_datetime(data_frame['DAT_OCOR_INFR'], format="%d/%m/%Y %H:%M")
        data_frame['DAT_EMIS_NOTF'] = pd.to_datetime(data_frame['DAT_EMIS_NOTF'], format="%d/%m/%Y")
        data_frame['DAT_LIMT_RECU'] = pd.to_datetime(data_frame['DAT_LIMT_RECU'], format="%d/%m/%Y")
        data_frame["VAL_INFR"] = data_frame['VAL_INFR'].map(lambda x: pd.to_numeric(str(x).replace(',', '.')))
        if 'DAT_CANC' in data_frame.columns and not bool(data_frame['DAT_CANC'].isnull().all()):
            data_frame['DAT_CANC'] = pd.to_datetime(data_frame['DAT_CANC'], format="%d/%m/%Y")
        data_frame = data_frame.drop(columns=['HORA'])
        return data_frame
    except Exception as e:
        raise ErrReadingFile(f"Erro ao processar o arquivo CSV. {e}", 500)


def parse_xls_infracoes(xls: Any) -> pd.DataFrame:
    """Processa o arquivo XLS de infrações retornando um DataFrame formatado"""
    try:
        data_frame = pd.read_excel(xls, header=0)
    except Exception as e:
        raise ErrReadingFile(f'Problema ao processar o arquivo no Load: {xls}. {e}', 500)

    try:
        data_frame['DAT_OCOR_INFR'] = data_frame['DAT_OCOR_INFR'].astype(str) + " " + data_frame['HORA'].astype(str)
        data_frame['DAT_OCOR_INFR'] = pd.to_datetime(data_frame['DAT_OCOR_INFR'], format="%Y-%m-%d %H:%M:%S")
        data_frame['DAT_EMIS_NOTF'] = pd.to_datetime(data_frame['DAT_EMIS_NOTF'], format="%Y-%m-%d")
        data_frame['DAT_LIMT_RECU'] = pd.to_datetime(data_frame['DAT_LIMT_RECU'], format="%Y-%m-%d")
        if 'DAT_CANC' in data_frame.columns and not bool(data_frame['DAT_CANC'].isnull().all()):
            data_frame['DAT_CANC'] = pd.to_datetime(data_frame['DAT_CANC'], format="%Y-%m-%d")

        data_frame = data_frame.drop(columns=['HORA'])
        data_frame.replace([np.nan], [None], inplace=True)
        return data_frame
    except Exception as e:
        raise ErrReadingFile(f'Problema ao corrigir datas e manipular colunas: {xls}. {e}', 500)
