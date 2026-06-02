import re
from typing import List, Dict, Any
from docx import Document
from classes.Recurso import Recurso
from classes.Conversores import Conversores
from repositories.database import get_db
from repositories.recurso_repository import RecursoRepository
from exceptions.CustomExceptions import (
    ErrDataPubli, ErrGetData, ErrInsertData, ErrNullInsert,
    ErrQuantityOfAtas, ErrIncorrectInstance
)

from handlers.log import logger



def getPrimeiraInstancia(date: Any, ata: Any) -> List[Dict[str, Any]]:
    """Retorna os recursos de primeira instância"""
    try:
        with get_db() as db:
            repo = RecursoRepository(db)
            return repo.get_primeira_instancia(date, ata)
    except Exception as e:
        logger.systemLog(e)
        raise ErrGetData("Erro ao recuperar os recursos de primeira instancia", 500)


def getSegundaInstancia(date: Any) -> List[Dict[str, Any]]:
    """Retorna os recursos de segunda instância"""
    try:
        with get_db() as db:
            repo = RecursoRepository(db)
            return repo.get_segunda_instancia(date)
    except Exception as e:
        logger.systemLog(e)
        raise ErrGetData("Erro ao recuperar os recursos de segunda instancia", 500)


def parseDocx(docx: Any, first_instance: bool = True) -> List[Dict[str, Any]]:
    """Realiza o parse do documento DOCX de resultado de recurso e extrai os KVP e retorna uma lista de recursos"""
    doc = Document(docx)
    recurso_primeira_instancia_list = []

    # Extraindo data da publicação
    data_publicacao_extracted = ""
    for paragraph in doc.paragraphs:
        data_publicacao_extracted += paragraph.text + " "

    padrao_data = r'PUBLICADO NO DI[ÁA]RIO OFICIAL DO MUNIC[ÍI]PIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})'
    match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)

    dat_publ = match_data_publicacao.group(1) if match_data_publicacao else None

    if dat_publ is not None:
        dat_publ = Conversores.converte_data(dat_publ)
    else:
        raise ErrDataPubli("Data de publicação não encontrada no documento", 400)

    # Extraindo Numero da ata
    num_atas = []
    padrao_num_ata = r'ATA\s+DA\s+(\d+)ª'
    for paragraph in doc.paragraphs:
        match_num_ata = re.search(padrao_num_ata, paragraph.text)
        if match_num_ata:
            num_atas.append(match_num_ata.group(1))

    qtd_atas = len(num_atas)
    qtd_tables = len(doc.tables)
    if (qtd_atas != qtd_tables) and first_instance:
        raise ErrQuantityOfAtas("Quantidade de atas encontradas difere da quantidade de tabelas", qtd_atas, qtd_tables, 400)

    for index, table in enumerate(doc.tables):
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            num_recurso = row_data[0]
            num_ai = normalizeAutoInfractionId(row_data[1])
            nom_conc = row_data[2]
            valida_resultado = str(row_data[3]).upper()
            if valida_resultado == 'IMPROCEDENTE':
                resultado = False
            else:
                resultado = True

            if not first_instance and num_atas[index] is not None:
                raise ErrIncorrectInstance("Instância incorreta. Importe como recurso de primeira instância")

            num_ata = num_atas[index] if first_instance else 0

            if num_recurso == 'RECURSO':
                continue

            recurso_primeira_instancia = Recurso(num_ata, num_recurso, num_ai, nom_conc, resultado, dat_publ)
            recurso_primeira_instancia_list.append(recurso_primeira_instancia.toDict())
    return recurso_primeira_instancia_list


def insertPrimeiraInstancia(recursos_primeira_instancia: List[Dict[str, Any]] | None) -> int:
    """Insere no banco de dados uma lista de recursos de primeira instância"""
    if recursos_primeira_instancia is None:
        raise ErrNullInsert('Lista de recursos vazia, nenhum registro inserido', 400)

    try:
        with get_db() as db:
            repo = RecursoRepository(db)
            count = repo.insert_primeira_instancia(recursos_primeira_instancia)
            return count
    except Exception as e:
        logger.systemLog(e)
        raise ErrInsertData("Erro ao inserir recursos de primeira instância no banco", 500)


def insertSegundaInstancia(recursos_segunda_instancia: List[Dict[str, Any]] | None) -> int:
    """Insere no banco de dados uma lista de recursos de segunda instância"""
    if recursos_segunda_instancia is None:
        raise ErrNullInsert('Lista de recursos vazia, nenhum registro inserido', 400)

    try:
        with get_db() as db:
            repo = RecursoRepository(db)
            count = repo.insert_segunda_instancia(recursos_segunda_instancia)
            return count
    except Exception as e:
        logger.systemLog(e)
        raise ErrInsertData("Erro ao inserir recursos de segunda instância no banco", 500)


def normalizeAutoInfractionId(ai: str) -> str:
    if not ai:
        return ai

    if ai[len(ai) - 1] == "-":
        return ai

    return f"{ai[:-1]}-{ai[-1]}"
