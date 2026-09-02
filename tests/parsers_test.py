import os
import tempfile

import numpy as np
import pandas as pd
import pytest
from docx import Document

from exceptions.CustomExceptions import ErrDataPubli, ErrInvalidFileData
from services.parsers.pyingestion.streams import (
    InfracoesTransformStream,
    RecursosDocxInputStream,
    normalize_auto_infraction_id,
)


def test_normalize_auto_infraction_id():
    assert normalize_auto_infraction_id("12345A") == "12345-A"
    assert normalize_auto_infraction_id("12345-A") == "12345-A"
    assert normalize_auto_infraction_id("1234-") == "1234-"
    assert normalize_auto_infraction_id("") == ""


def test_infracoes_transform_accepts_brazilian_currency_and_missing_value():
    data_frame = pd.DataFrame(
        {
            "NUM_AI": ["12345-A", "12346-A"],
            "DAT_OCOR_INFR": ["01/09/2026", "02/09/2026"],
            "HORA": ["10:30", "11:45"],
            "DAT_EMIS_NOTF": ["01/09/2026", np.nan],
            "DAT_LIMT_RECU": ["15/09/2026", "16/09/2026"],
            "VAL_INFR": ["R$ 1.498,91", np.nan],
        }
    )
    stream = InfracoesTransformStream(
        datetime_format="%d/%m/%Y %H:%M",
        date_format="%d/%m/%Y",
        convert_val_infr=True,
    )

    records = stream.transform(data_frame)

    assert records[0]["VAL_INFR"] == 1498.91
    assert records[1]["VAL_INFR"] is None
    assert records[1]["DAT_EMIS_NOTF"] is None
    assert records[1]["DAT_LIMT_RECU"] == pd.Timestamp("2026-09-16")


def test_recursos_docx_stream():
    # Create a temporary DOCX file with expected structure
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc_path = temp_file.name

    try:
        doc = Document()
        # Add publication date paragraph
        doc.add_paragraph(
            "PUBLICADO NO DIARIO OFICIAL DO MUNICIPIO DE BELO HORIZONTE EM 15/05/2026"
        )
        # Add Session/Ata header paragraph
        doc.add_paragraph("ATA DA 5ª SESSÃO ORDINÁRIA")

        # Add table
        table = doc.add_table(rows=2, cols=4)
        # Row 0: Headers
        table.cell(0, 0).text = "RECURSO"
        table.cell(0, 1).text = "AUTO DE INFRAÇÃO"
        table.cell(0, 2).text = "RECORRENTE"
        table.cell(0, 3).text = "DECISÃO"
        # Row 1: Valid Data
        table.cell(1, 0).text = "123/2026"
        table.cell(1, 1).text = "345678A"
        table.cell(1, 2).text = "Consórcio BH Leste"
        table.cell(1, 3).text = "IMPROCEDENTE"

        doc.save(doc_path)

        # Parse the docx using the new InputStream
        stream = RecursosDocxInputStream(first_instance=True)
        recursos = list(stream.read(doc_path))
        assert len(recursos) == 1
        recurso = recursos[0]
        assert recurso["NUM_RECURSO"] == "123/2026"
        assert recurso["NUM_ATA"] == "5"
        assert recurso["NUM_AI"] == "345678-A"
        assert recurso["NOM_CONC"] == "Consórcio BH Leste"
        assert recurso["RESULTADO"] is False
        assert recurso["DAT_PUBL"] == "2026-05-15"

    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_recursos_docx_stream_missing_date():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc_path = temp_file.name

    try:
        doc = Document()
        doc.add_paragraph("ATA DA 5ª SESSÃO ORDINÁRIA")
        doc.save(doc_path)

        with pytest.raises(ErrDataPubli):
            stream = RecursosDocxInputStream(first_instance=True)
            list(stream.read(doc_path))

    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_infracoes_transform_missing_ai():
    data_frame = pd.DataFrame(
        {
            "NUM_AI": ["", "12346-A"],
            "DAT_OCOR_INFR": ["01/09/2026", "02/09/2026"],
            "DAT_LIMT_RECU": ["15/09/2026", "16/09/2026"],
        }
    )
    stream = InfracoesTransformStream(
        datetime_format="%d/%m/%Y %H:%M",
        date_format="%d/%m/%Y",
        convert_val_infr=True,
    )
    with pytest.raises(ErrInvalidFileData) as excinfo:
        stream.transform(data_frame)
    assert "NUM_AI" in excinfo.value.friendly_message


def test_recursos_docx_stream_invalid_columns():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc_path = temp_file.name

    try:
        doc = Document()
        doc.add_paragraph(
            "PUBLICADO NO DIARIO OFICIAL DO MUNICIPIO DE BELO HORIZONTE EM 15/05/2026"
        )
        doc.add_paragraph("ATA DA 5ª SESSÃO ORDINÁRIA")
        table = doc.add_table(rows=2, cols=3)  # Only 3 cols, should be 4
        table.cell(0, 0).text = "RECURSO"
        table.cell(1, 0).text = "123/2026"
        table.cell(1, 1).text = "345678A"
        doc.save(doc_path)

        with pytest.raises(ErrInvalidFileData) as excinfo:
            stream = RecursosDocxInputStream(first_instance=True)
            list(stream.read(doc_path))
        assert "4 eram esperadas" in excinfo.value.friendly_message

    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)
